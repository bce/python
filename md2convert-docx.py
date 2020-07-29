#!/usr/bin/python3
import os
import docx
from docx.shared import RGBColor
def readfile(path,wordpath):
    doc = docx.Document()
    for file in os.listdir(path):
        if os.path.isdir(os.path.join(path,file)):
            flist = os.listdir(os.path.join(path,file))
            p=doc.add_paragraph()
            r=p.add_run(file+'\n')
            r.font.bold = True
            # run.font.color.rgb=RGBColor((0,0,255))
            flist.sort(key=lambda x: x.split('.')[1])
            for ff in flist:
                fmd = open(os.path.join(path,file,ff), 'r',encoding='utf-8')
                #
                if ff.endswith('.md'):   #��.md��β���ļ�
                    try :
                        doc.add_paragraph(fmd.readlines())
                    except ValueError:
                        pass
                else:
                    rr=doc.add_paragraph()
                    for line in fmd.readlines():
                        rr.add_run(line)
                doc.add_paragraph('')
            print(flist)
    doc.save(wordpath)
            #
                # print(os.path.join(path,file,f1))
readfile('./file','./file.docx')    #��file�ļ����������ļ�д��һ��word�ļ���